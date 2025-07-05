import streamlit as st
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import io

class CVTemplateGenerator:
    """Generates professional CV templates for different roles"""
    
    def __init__(self):
        self.templates = {
            'student': self.create_student_cv,
            'entrepreneur': self.create_entrepreneur_cv,
            'researcher': self.create_researcher_cv,
            'artist': self.create_artist_cv,
            'nonprofit': self.create_nonprofit_cv
        }
    
    def generate_cv(self, profile_data, role):
        """Generate CV based on role and profile data"""
        template_func = self.templates.get(role, self.create_student_cv)
        return template_func(profile_data)
    
    def create_student_cv(self, profile_data):
        """Create a student-focused CV template"""
        doc = Document()
        
        # Header
        header = doc.add_heading(profile_data.get('name', 'Your Name'), 0)
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Contact Info
        contact = doc.add_paragraph()
        contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact.add_run(f"📧 {profile_data.get('email', 'your.email@example.com')} | ")
        contact.add_run(f"📱 {profile_data.get('phone', 'Your Phone Number')}")
        
        # Education Section
        doc.add_heading('🎓 Education', level=1)
        edu_para = doc.add_paragraph()
        edu_para.add_run(f"Major: {profile_data.get('major', 'Your Major')}\n").bold = True
        edu_para.add_run(f"Expected Graduation: {profile_data.get('graduation_year', '2025')}\n")
        if profile_data.get('gpa'):
            edu_para.add_run(f"GPA: {profile_data.get('gpa')}/4.0\n")
        edu_para.add_run(f"Relevant Coursework: {profile_data.get('coursework', 'List your relevant courses')}")
        
        # Projects Section
        if profile_data.get('projects'):
            doc.add_heading('💻 Projects', level=1)
            doc.add_paragraph(profile_data.get('projects'))
        
        # Experience Section
        if profile_data.get('internships'):
            doc.add_heading('💼 Experience', level=1)
            doc.add_paragraph(profile_data.get('internships'))
        
        # Activities Section
        if profile_data.get('extracurricular'):
            doc.add_heading('🏆 Activities & Leadership', level=1)
            doc.add_paragraph(profile_data.get('extracurricular'))
        
        # Awards Section
        if profile_data.get('awards'):
            doc.add_heading('🏅 Awards & Recognition', level=1)
            doc.add_paragraph(profile_data.get('awards'))
        
        return doc
    
    def create_entrepreneur_cv(self, profile_data):
        """Create an entrepreneur-focused CV template"""
        doc = Document()
        
        # Header
        header = doc.add_heading(profile_data.get('name', 'Your Name'), 0)
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Contact Info
        contact = doc.add_paragraph()
        contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact.add_run(f"📧 {profile_data.get('email', 'your.email@example.com')} | ")
        contact.add_run(f"📱 {profile_data.get('phone', 'Your Phone Number')}")
        
        # Executive Summary
        doc.add_heading('💼 Executive Summary', level=1)
        summary = doc.add_paragraph()
        summary.add_run(f"Entrepreneur with expertise in {profile_data.get('market_focus', 'your market focus')}. ")
        summary.add_run(f"Founded {profile_data.get('ventures', 'describe your ventures')}.")
        
        # Business Experience
        doc.add_heading('🚀 Ventures & Business Experience', level=1)
        doc.add_paragraph(profile_data.get('business_experience', 'Describe your business experience'))
        
        # Key Achievements
        if profile_data.get('achievements'):
            doc.add_heading('🏆 Key Achievements', level=1)
            doc.add_paragraph(profile_data.get('achievements'))
        
        # Skills
        doc.add_heading('💡 Core Skills', level=1)
        doc.add_paragraph(profile_data.get('skills', 'List your key business skills'))
        
        return doc
    
    def create_researcher_cv(self, profile_data):
        """Create a researcher-focused CV template"""
        doc = Document()
        
        # Header
        header = doc.add_heading(profile_data.get('name', 'Your Name'), 0)
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Contact Info
        contact = doc.add_paragraph()
        contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact.add_run(f"📧 {profile_data.get('email', 'your.email@example.com')} | ")
        contact.add_run(f"📱 {profile_data.get('phone', 'Your Phone Number')}")
        
        # Research Interests
        doc.add_heading('🔬 Research Interests', level=1)
        doc.add_paragraph(profile_data.get('research_area', 'Describe your research area and interests'))
        
        # Education
        doc.add_heading('🎓 Education', level=1)
        doc.add_paragraph(profile_data.get('education', 'List your educational background'))
        
        # Publications
        if profile_data.get('publications'):
            doc.add_heading('📚 Publications', level=1)
            doc.add_paragraph(profile_data.get('publications'))
        
        # Grants & Funding
        if profile_data.get('grants'):
            doc.add_heading('💰 Grants & Funding', level=1)
            doc.add_paragraph(profile_data.get('grants'))
        
        # Conferences
        if profile_data.get('conferences'):
            doc.add_heading('🎤 Conferences & Presentations', level=1)
            doc.add_paragraph(profile_data.get('conferences'))
        
        return doc
    
    def create_artist_cv(self, profile_data):
        """Create an artist-focused CV template"""
        doc = Document()
        
        # Header
        header = doc.add_heading(profile_data.get('name', 'Your Name'), 0)
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Contact Info
        contact = doc.add_paragraph()
        contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact.add_run(f"📧 {profile_data.get('email', 'your.email@example.com')} | ")
        contact.add_run(f"🎨 {profile_data.get('portfolio_url', 'your-portfolio.com')}")
        
        # Artist Statement
        doc.add_heading('🎨 Artist Statement', level=1)
        statement = doc.add_paragraph()
        statement.add_run(f"Working primarily in {profile_data.get('medium', 'your medium')}, ")
        statement.add_run(f"my style is characterized by {profile_data.get('style', 'describe your style')}. ")
        statement.add_run(f"My work is inspired by {profile_data.get('inspiration', 'your inspiration')}.")
        
        # Exhibitions
        if profile_data.get('exhibitions'):
            doc.add_heading('🖼️ Exhibitions', level=1)
            doc.add_paragraph(profile_data.get('exhibitions'))
        
        # Awards
        if profile_data.get('awards'):
            doc.add_heading('🏆 Awards & Recognition', level=1)
            doc.add_paragraph(profile_data.get('awards'))
        
        # Collections
        if profile_data.get('collections'):
            doc.add_heading('🏛️ Collections', level=1)
            doc.add_paragraph(profile_data.get('collections'))
        
        return doc
    
    def create_nonprofit_cv(self, profile_data):
        """Create a nonprofit-focused CV template"""
        doc = Document()
        
        # Header
        header = doc.add_heading(profile_data.get('name', 'Your Name'), 0)
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Contact Info
        contact = doc.add_paragraph()
        contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact.add_run(f"📧 {profile_data.get('email', 'your.email@example.com')} | ")
        contact.add_run(f"📱 {profile_data.get('phone', 'Your Phone Number')}")
        
        # Mission Statement
        doc.add_heading('🎯 Mission & Impact', level=1)
        mission = doc.add_paragraph()
        mission.add_run(f"Organization: {profile_data.get('organization', 'Your Organization')}\n").bold = True
        mission.add_run(f"Mission: {profile_data.get('mission', 'Describe your mission')}\n")
        mission.add_run(f"Impact Area: {profile_data.get('impact_area', 'Your impact area')}")
        
        # Programs
        if profile_data.get('programs'):
            doc.add_heading('📋 Programs & Initiatives', level=1)
            doc.add_paragraph(profile_data.get('programs'))
        
        # Partnerships
        if profile_data.get('partnerships'):
            doc.add_heading('🤝 Partnerships', level=1)
            doc.add_paragraph(profile_data.get('partnerships'))
        
        # Outcomes
        if profile_data.get('outcomes'):
            doc.add_heading('📊 Impact & Outcomes', level=1)
            doc.add_paragraph(profile_data.get('outcomes'))
        
        return doc
    
    def save_cv_to_bytes(self, doc):
        """Save document to bytes for download"""
        bio = io.BytesIO()
        doc.save(bio)
        bio.seek(0)
        return bio
